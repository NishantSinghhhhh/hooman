import os
import time
import uuid
import shutil
import json
from typing import Dict, Any, Optional
from crewai import Agent, Task, Crew
import openai
from pathlib import Path
import PyPDF2
import docx

# Import the Pixeltable insert and query functions
try:
    from src.ingestor.ingestor import (
        insert_document_record,
        insert_tracking_record,
        get_user_records,
        get_token_usage_summary
    )
    PIXELTABLE_AVAILABLE = True
    print("✅ Pixeltable integration loaded for DocumentAgent")
except ImportError:
    print("⚠️ Pixeltable integration not available - continuing without database storage")
    PIXELTABLE_AVAILABLE = False

class DocumentAgent:
    """Specialized agent for document processing and analysis."""

    def __init__(self):
        """Initialize the Document Agent with OpenAI client and data directory."""
        self.openai_client = openai.OpenAI(
            api_key=os.getenv("OPENAI_API_KEY")
        )
        
        # Create data directory if it doesn't exist
        self.data_dir = os.path.join(os.getcwd(), "data", "documents")
        os.makedirs(self.data_dir, exist_ok=True)
        print(f"📁 Data directory for documents: {self.data_dir}")
        print("✅ DocumentAgent initialized")

    def save_document_to_data_folder(self, source_path: str, user_id: str = "anonymous") -> str:
        """Saves the document to the data folder with a unique name."""
        try:
            original_filename = os.path.basename(source_path)
            name, ext = os.path.splitext(original_filename)
            timestamp = int(time.time())
            unique_id = str(uuid.uuid4())[:8]
            new_filename = f"{user_id}_{timestamp}_{unique_id}_{name}{ext}"
            new_path = os.path.join(self.data_dir, new_filename)
            shutil.copy2(source_path, new_path)
            print(f"💾 Document saved to data folder: {new_filename}")
            return new_path
        except Exception as e:
            print(f"❌ Error saving document to data folder: {str(e)}")
            return source_path

    def cleanup_temp_file(self, file_path: str) -> None:
        """Cleans up a temporary file if it's not in the data folder."""
        try:
            if not file_path.startswith(self.data_dir) and os.path.exists(file_path):
                os.remove(file_path)
                print(f"🗑️ Cleaned up temporary file: {os.path.basename(file_path)}")
        except Exception as e:
            print(f"⚠️ Could not clean up temporary file: {str(e)}")
            
    def get_document_info(self, doc_path: str) -> Dict[str, Any]:
        """Gets detailed information about the document file."""
        try:
            file_stats = os.stat(doc_path)
            file_ext = Path(doc_path).suffix.lower()
            page_count = 0
            
            if file_ext == '.pdf':
                # Use strict=False to handle a wider variety of PDFs
                with open(doc_path, 'rb') as f:
                    pdf = PyPDF2.PdfReader(f, strict=False)
                    page_count = len(pdf.pages)
            elif file_ext == '.docx':
                doc = docx.Document(doc_path)
                # Estimate pages based on word count (approx. 300 words/page)
                word_count = sum(len(p.text.split()) for p in doc.paragraphs)
                page_count = max(1, (word_count + 299) // 300)
            else: # .txt or other
                page_count = 1

            return {
                "filename": os.path.basename(doc_path),
                "file_size_bytes": file_stats.st_size,
                "document_type": file_ext.replace('.', ''),
                "page_count": page_count,
                "modified_time": file_stats.st_mtime,
            }
        except Exception as e:
            print(f"❌ Error getting document info: {str(e)}")
            return {"error": str(e)}

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extracts text from a PDF file."""
        text = ""
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file, strict=False)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        return text.strip()

    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extracts text from a DOCX file."""
        doc = docx.Document(docx_path)
        return "\n".join([p.text for p in doc.paragraphs]).strip()

    def extract_text_from_txt(self, txt_path: str) -> str:
        """Extracts text from a TXT file, trying multiple encodings."""
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            with open(txt_path, 'r', encoding='latin-1') as file:
                return file.read().strip()
        except Exception as e:
            print(f"❌ Error reading TXT file: {e}")
            return ""

    def extract_document_text(self, document_path: str) -> Dict[str, Any]:
        """Extracts text from various document formats."""
        try:
            file_ext = Path(document_path).suffix.lower()
            print(f"📄 Extracting text from {os.path.basename(document_path)}")
            
            text_extractors = {
                '.pdf': self.extract_text_from_pdf,
                '.docx': self.extract_text_from_docx,
                '.txt': self.extract_text_from_txt
            }
            
            if file_ext not in text_extractors:
                return {"success": False, "error": f"Unsupported file format: {file_ext}"}

            text = text_extractors[file_ext](document_path)

            if not text:
                return {"success": False, "error": "No text could be extracted."}

            return {
                "success": True,
                "text": text,
                "word_count": len(text.split()),
                "char_count": len(text)
            }
        except Exception as e:
            return {"success": False, "error": f"Text extraction failed: {str(e)}"}

    def analyze_document_with_openai(self, doc_text: str, query: str) -> Dict[str, Any]:
        """Analyzes document text using the OpenAI API."""
        try:
            print("🔍 Analyzing document with OpenAI API...")
            if not query.strip():
                prompt = "Provide a concise summary of the following document. Identify the main topics, key arguments, and conclusions."
            else:
                prompt = f"Based on the document provided, please answer the following question: '{query}'"

            # Use a larger context window for documents
            response = self.openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant that analyzes documents."},
                    {"role": "user", "content": f"{prompt}\n\n--- DOCUMENT TEXT ---\n{doc_text[:12000]}"},
                ],
                max_tokens=1500,
                temperature=0.2,
            )
            analysis = response.choices[0].message.content
            tokens_used = response.usage.total_tokens if response.usage else 0
            
            return {
                "success": True,
                "analysis": analysis,
                "model_used": "gpt-4o",
                "tokens_used": tokens_used,
            }
        except Exception as e:
            return {"success": False, "error": f"OpenAI API error: {str(e)}", "tokens_used": 0}

    def create_document_analysis_agent(self) -> Agent:
        """Creates a specialized document analysis agent."""
        return Agent(
            role='Senior Document Analyst',
            goal='Diligently analyze document text to extract key insights and answer user queries.',
            backstory='An expert in semantic analysis and information retrieval, you can dissect any document to find the most relevant and crucial information.',
            verbose=True,
            allow_delegation=False
        )

    def create_content_synthesizer_agent(self) -> Agent:
        """Creates an agent to synthesize and format the final response."""
        return Agent(
            role='Expert Content Synthesizer',
            goal='Transform raw analysis into a clear, structured, and user-friendly response.',
            backstory='A master of communication, you specialize in organizing complex information into easily digestible summaries, reports, and answers.',
            verbose=True,
            allow_delegation=False
        )
            
    def store_to_pixeltable(self, doc_path: str, query: str, result: Dict[str, Any], user_id: str, processing_time: float, success: bool) -> bool:
        """Stores the document processing results to Pixeltable."""
        if not PIXELTABLE_AVAILABLE:
            print("⚠️ Pixeltable not available - skipping database storage.")
            return False
            
        try:
            print("💾 Storing document results to Pixeltable database...")
            doc_info = self.get_document_info(doc_path)
            tokens_used = result.get('tokens', 0)  # Get tokens from top level
            
            # Insert document record only on success
            if success:
                res_data = result.get('result', {})
                crewai_result = {
                    "primary_analysis": res_data.get('analysis', ''),
                    "enhanced_response": res_data.get('enhanced_response', ''),
                    "status": "success"
                }
                insert_document_record(
                    user_id=user_id,
                    document_path=doc_path,
                    document_type=doc_info.get('document_type', ''),
                    page_count=doc_info.get('page_count', 0),
                    query=query,
                    crewai_result=crewai_result,
                    tokens_used=int(tokens_used),
                    context=f"Document analysis - {result.get('processing_method', 'unknown')}",
                    metadata=doc_info
                )
            
            # Insert tracking record for both success and failure
            insert_tracking_record(
                user_id=user_id,
                agent_type='document',
                table_name='demo.documents',
                record_id=f"doc_{int(time.time())}",
                query=query,
                tokens_used=int(tokens_used),
                processing_time=processing_time,
                success=success,
                error_message=result.get('error')
            )
            print("✅ Database storage completed successfully.")
            return True
                
        except Exception as e:
            print(f"❌ Database storage error: {str(e)}")
            return False

    def quick_analyze(self, document_path: str, query: str = "", user_id: str = "anonymous") -> Dict[str, Any]:
        """Quick analysis using OpenAI only (faster, less detailed)."""
        start_time = time.time()
        original_path = document_path
        saved_doc_path = None
        try:
            print("⚡ Starting quick document analysis...")
            saved_doc_path = self.save_document_to_data_folder(document_path, user_id)
            
            extraction_result = self.extract_document_text(saved_doc_path)
            if not extraction_result.get("success"):
                raise ValueError(extraction_result.get("error"))

            document_text = extraction_result.get("text", "")
            
            # Analyze with OpenAI only
            openai_result = self.analyze_document_with_openai(document_text, query)
            if not openai_result.get("success"):
                raise ValueError(openai_result.get("error"))
            
            processing_time = time.time() - start_time
            total_tokens = openai_result.get('tokens_used', 0)
            
            final_result = {
                'success': True,
                'tokens': total_tokens,  # TOP LEVEL TOKEN COUNT
                'result': {
                    "analysis": openai_result.get("analysis", ""),
                    "technical_details": {
                        "tokens_used": total_tokens,
                        "model_used": openai_result.get("model_used", "gpt-4o"),
                        "word_count": extraction_result.get("word_count", 0),
                        "char_count": extraction_result.get("char_count", 0)
                    },
                },
                'query': query,
                'processing_method': 'text_extraction + openai_only',
                'processing_time': processing_time,
            }
            
            self.store_to_pixeltable(saved_doc_path, query, final_result, user_id, processing_time, True)
            print(f"📊 Total tokens used: {total_tokens}")
            print(f"✅ Quick document analysis completed in {processing_time:.2f}s")
            return final_result
            
        except Exception as e:
            processing_time = time.time() - start_time
            print(f"❌ Quick document analysis failed: {str(e)}")
            error_result = {
                'success': False, 
                'tokens': 0,  # No tokens on error
                'error': str(e), 
                'query': query, 
                'processing_time': processing_time
            }
            if saved_doc_path:
                self.store_to_pixeltable(saved_doc_path, query, error_result, user_id, processing_time, False)
            return error_result
        finally:
            if original_path != saved_doc_path and saved_doc_path:
                self.cleanup_temp_file(original_path)

    def process_document(self, document_path: str, query: str = "", user_id: str = "anonymous") -> Dict[str, Any]:
        """Full, in-depth analysis using text extraction and CrewAI."""
        start_time = time.time()
        original_path = document_path
        saved_doc_path = None
        try:
            print("⚙️ Starting full document processing...")
            saved_doc_path = self.save_document_to_data_folder(document_path, user_id)
            
            extraction_result = self.extract_document_text(saved_doc_path)
            if not extraction_result.get("success"):
                raise ValueError(extraction_result.get("error"))

            document_text = extraction_result.get("text", "")
            
            # First get OpenAI analysis for token counting
            openai_result = self.analyze_document_with_openai(document_text, query)
            openai_tokens = openai_result.get('tokens_used', 0) if openai_result.get('success') else 0
            
            # CrewAI setup
            analysis_agent = self.create_document_analysis_agent()
            synthesizer_agent = self.create_content_synthesizer_agent()
            
            analysis_task = Task(
                description=f"Analyze the following document text based on the user's query.\n\nQuery: {query}\n\nText: {document_text}", 
                expected_output="A detailed analysis addressing the query.", 
                agent=analysis_agent
            )
            synthesis_task = Task(
                description="Synthesize the analysis from the previous task into a final, user-friendly response.", 
                expected_output="A polished, comprehensive answer.", 
                agent=synthesizer_agent,
                context=[analysis_task]
            )
            
            crew = Crew(agents=[analysis_agent, synthesizer_agent], tasks=[analysis_task, synthesis_task], verbose=2)
            crew_output = crew.kickoff()
            
            # Get crew tokens
            crew_tokens = getattr(crew.usage_metrics, 'total_tokens', 0) if hasattr(crew, 'usage_metrics') else 0
            total_tokens = openai_tokens + crew_tokens
            
            processing_time = time.time() - start_time
            final_result = {
                'success': True,
                'tokens': total_tokens,  # TOP LEVEL TOKEN COUNT
                'result': {
                    "analysis": openai_result.get("analysis", "") if openai_result.get('success') else "",
                    "enhanced_response": str(crew_output),
                    "technical_details": {
                        "tokens_used": total_tokens,
                        "token_breakdown": {
                            "openai_analysis": openai_tokens,
                            "crew_enhancement": crew_tokens
                        },
                        "word_count": extraction_result.get("word_count", 0),
                        "char_count": extraction_result.get("char_count", 0)
                    },
                },
                'query': query,
                'processing_method': 'text_extraction + openai + crewai_analysis',
                'processing_time': processing_time,
            }
            
            self.store_to_pixeltable(saved_doc_path, query, final_result, user_id, processing_time, True)
            print(f"📊 Total tokens used: {total_tokens}")
            print(f"✅ Full document processing completed in {processing_time:.2f}s")
            return final_result
            
        except Exception as e:
            processing_time = time.time() - start_time
            print(f"❌ Full document processing failed: {str(e)}")
            error_result = {
                'success': False, 
                'tokens': 0,  # No tokens on error
                'error': str(e), 
                'query': query, 
                'processing_time': processing_time
            }
            if saved_doc_path:
                self.store_to_pixeltable(saved_doc_path, query, error_result, user_id, processing_time, False)
            return error_result
        finally:
            # Cleanup the original temp file if it was copied
            if original_path != saved_doc_path and saved_doc_path:
                self.cleanup_temp_file(original_path)

# Example usage
if __name__ == "__main__":
    agent = DocumentAgent()

    # Create a dummy text file for testing
    dummy_doc_path = "/tmp/test_document.txt"
    with open(dummy_doc_path, "w", encoding='utf-8') as f:
        f.write("This is a test document about two AI frameworks.\n")
        f.write("Pixeltable is a library for managing and analyzing large-scale data, especially for AI.\n")
        f.write("CrewAI is a framework designed to orchestrate autonomous AI agents, making them work together on complex tasks.\n")
    print(f"\nCreated a dummy test document at: {dummy_doc_path}")

    # TEST QUICK ANALYSIS
    print("\n" + "="*20 + " TESTING QUICK ANALYSIS " + "="*20)
    quick_result = agent.quick_analyze(
        document_path=dummy_doc_path,
        query="What is Pixeltable?",
        user_id="doc_user_quick_789"
    )
    print(f"🔢 Tokens used in quick analysis: {quick_result.get('tokens', 0)}")

    # TEST FULL ANALYSIS
    print("\n" + "="*20 + " TESTING FULL ANALYSIS " + "="*20)
    full_result = agent.process_document(
        document_path=dummy_doc_path,
        query="Compare Pixeltable and CrewAI based on the text provided.",
        user_id="doc_user_full_999"
    )
    print(f"🔢 Tokens used in full analysis: {full_result.get('tokens', 0)}")

    if full_result.get('success'):
        print("Enhanced Response:", full_result.get('result', {}).get('enhanced_response'))

    # TEST THE QUERY FUNCTIONS
    if PIXELTABLE_AVAILABLE:
        print("\n" + "="*20 + " PULLING DATA FROM PIXELTABLE " + "="*20)

        print("\n--- Overall Token Usage Summary ---")
        summary = get_token_usage_summary()
        print(json.dumps(summary, indent=2))
        
        print("\n--- Document Records for 'doc_user_quick_789' ---")
        user_records = get_user_records(user_id="doc_user_quick_789", agent_type="document")
        if 'document' in user_records and hasattr(user_records['document'], 'head'):
             print(user_records['document'].head())
        else:
             print(user_records)
    
    # Clean up the dummy file
    agent.cleanup_temp_file(dummy_doc_path)